from __future__ import annotations
from typing import Dict, Any
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from .utils import SECTION_ORDER


def to_docx(plan: Dict[str, Any], path: str) -> str:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    doc.add_heading(plan.get("title", "Lesson Plan"), level=1)
    sections = plan.get("sections", {})

    for section in SECTION_ORDER:
        if section in sections and sections[section]:
            doc.add_heading(section, level=2)
            content = sections[section]
            if isinstance(content, list):
                for item in content:
                    doc.add_paragraph(str(item), style="List Bullet")
            else:
                doc.add_paragraph(str(content))

    doc.save(path)
    return path


def to_pdf(plan: Dict[str, Any], path: str) -> str:
    c = canvas.Canvas(path, pagesize=A4)
    width, height = A4
    x, y = 40, height - 40

    def write_line(text: str):
        nonlocal y
        if y < 60:
            c.showPage()
            y = height - 40
        c.drawString(x, y, text)
        y -= 16

    title = plan.get("title", "Lesson Plan")
    c.setFont("Helvetica-Bold", 16)
    write_line(title)
    c.setFont("Helvetica", 11)

    sections = plan.get("sections", {})
    for section in SECTION_ORDER:
        content = sections.get(section)
        if not content:
            continue
        c.setFont("Helvetica-Bold", 12)
        write_line(section)
        c.setFont("Helvetica", 11)
        if isinstance(content, list):
            for item in content:
                write_line(f"• {str(item)}")
        else:
            for line in str(content).splitlines():
                write_line(line)
        y -= 4

    c.save()
    return path